from docx import Document
from datetime import date
import json
from typing import Dict

def genera_report_docx(cliente: Dict, profilo: Dict, path_output: str) -> str:
    doc = Document()

    doc.add_heading("Profiling Rischio AML", 0)
    doc.add_paragraph(f"Data generazione: {date.today().isoformat()}")

    doc.add_heading("Dati Cliente", level=1)
    doc.add_paragraph(f"Nome: {cliente['nome']}")
    doc.add_paragraph(f"Cognome: {cliente['cognome']}")
    doc.add_paragraph(f"Nascita: {cliente['data_nascita']}")
    doc.add_paragraph(f"Nazionalità: {cliente['nazionalità']}")

    doc.add_heading("Profilazione Rischio", level=1)
    doc.add_paragraph(f"Punteggio: {profilo['punteggio']}")
    doc.add_paragraph(f"Livello Rischio: {profilo['livello']}")
    doc.add_paragraph(f"Articoli Analizzati: {profilo['articoli_analizzati']}")

    doc.add_heading("Fonti Online", level=1)
    fonti = json.loads(profilo['fonti'])
    for url, score in fonti.items():
        doc.add_paragraph(f"- {url} (punteggio: {score})")

    output_path = f"{path_output}/report_aml_{cliente['id']}.docx"
    doc.save(output_path)
    return output_path